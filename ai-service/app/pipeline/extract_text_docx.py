"""
extraction/docx_extractor.py

مسؤول فقط عن استخراج النص الخام من ملفات DOCX.
منقول من الـ Kaggle Notebook (extract_text_from_cv - جزء الـ DOCX) مع تحسينات:
- بيستخرج نص من الجداول (tables) كمان، مش بس الـ paragraphs العادية
  (كتير من الـ CVs بتستخدم جداول لتنظيم الـ layout، والـ Notebook الأصلي كان بيتجاهلها)
- error handling أوضح للملفات التالفة
"""

import os
import docx  # python-docx


class DOCXExtractionError(Exception):
    """بترفع لو حصلت مشكلة أثناء قراءة ملف DOCX (ملف تالف، صيغة غير متوقعة، إلخ)."""
    pass


def extract_text_from_docx(file_path: str) -> str:
    """
    يستخرج النص الخام من ملف DOCX باستخدام python-docx.
    بيغطي الـ paragraphs العادية + النصوص جوه الجداول (tables).

    Args:
        file_path: المسار الكامل لملف الـ DOCX.

    Returns:
        النص المستخرج من الملف كامل.

    Raises:
        FileNotFoundError: لو الملف مش موجود أصلاً.
        DOCXExtractionError: لو الملف تالف أو مش قادرين نقرأه لأي سبب.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"الملف مش موجود: {file_path}")

    try:
        document = docx.Document(file_path)
    except Exception as e:
        raise DOCXExtractionError(f"تعذر قراءة ملف الـ DOCX: {e}") from e

    text_parts = []

    # 1. النص من الـ paragraphs العادية (زي الـ Notebook الأصلي)
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # 2. النص من جوه الجداول (إضافة جديدة - كتير من الـ CVs بتستخدم جداول)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    text_parts.append(cell_text)

    full_text = "\n".join(text_parts).strip()

    if not full_text:
        raise DOCXExtractionError(
            "لم يتم استخراج أي نص من ملف الـ DOCX. قد يكون الملف فارغًا."
        )

    return full_text